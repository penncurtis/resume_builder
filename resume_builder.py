import io
from typing import Dict, Any
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from templates import get_template_styles

class ResumeBuilder:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def generate_preview(self, resume_data: Dict, template: str) -> str:
        """Generate plain text preview of resume"""
        template_styles = get_template_styles(template)
        
        # Build contact info
        contact_parts = []
        if resume_data['contact']['email']:
            contact_parts.append(resume_data['contact']['email'])
        if resume_data['contact']['phone']:
            contact_parts.append(resume_data['contact']['phone'])
        if resume_data['contact']['location']:
            contact_parts.append(resume_data['contact']['location'])
        if resume_data['contact']['linkedin']:
            contact_parts.append(f"LinkedIn: {resume_data['contact']['linkedin']}")
        if resume_data['contact']['github']:
            contact_parts.append(f"GitHub: {resume_data['contact']['github']}")
        
        # Use Lorem ipsum placeholders if no real data
        if contact_parts:
            contact_info = " | ".join(contact_parts)
        else:
            contact_info = "Lorem City, LC | 555-555-5555 | lorem@ipsum.com \nLinkedIn | Github | Website"
        
        # Build experience section
        experience_text = ""
        if resume_data['experience']:
            for exp in resume_data['experience']:
                date_range = ""
                if exp.get('start_date') and exp.get('end_date'):
                    date_range = f" ({exp['start_date']} - {exp['end_date']})"
                elif exp.get('start_date'):
                    date_range = f" ({exp['start_date']} - Present)"
                
                experience_text += f"""
{exp.get('title', 'Job Title')} — {exp.get('company', 'Company Name')}{date_range}
{exp.get('description', 'Job description will appear here')}

"""
        else:
            # Lorem ipsum placeholder for experience - matching exact format
            experience_text = """
Role | Company Name
| City, State
Website 00/0000 - 00/0000
Lorem ipsum dolor sit amet, consectetur adipiscing elit.
Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.
Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.
Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.
Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.
Project Name – Platform / Description
Lorem ipsum dolor sit amet, consectetur adipiscing elit.
Integer vel sapien ut libero venenatis aliquet.
Curabitur rutrum, massa at sodales cursus, orci augue convallis nulla, ut faucibus lectus metus nec eros.

Role | Company Name
| City, State
00/0000 - 00/0000
Lorem ipsum dolor sit amet, consectetur adipiscing elit.
Vivamus interdum, nunc in dapibus volutpat, nunc ex commodo lectus, id condimentum massa ante sit amet nisl.
Pellentesque habitant morbi tristique senectus et netus et malesuada fames ac turpis egestas.
Sed tristique sapien eget sem ultricies, sit amet aliquam ipsum ultricies.
Fusce vel nunc eget sapien feugiat ullamcorper.
"""
        
        # Build education section
        education_text = ""
        if resume_data['education']:
            for edu in resume_data['education']:
                year_str = f" ({edu.get('year', '')})" if edu.get('year') else ""
                education_text += f"{edu.get('degree', 'Degree')} — {edu.get('school', 'School Name')}{year_str}\n"
        else:
            # Lorem ipsum placeholder for education - matching exact format
            education_text = """School Name | City, State
Program Name 00/0000 - 00/0000
University Name – School or Program | City, State
Degree Name 00/0000 - 00/0000
Lorem ipsum dolor sit amet, consectetur adipiscing elit.
Integer nec odio. Praesent libero. Sed cursus ante dapibus diam.

"""
        
        # Build skills section
        if resume_data['skills']:
            skills_text = ", ".join(resume_data['skills'])
        else:
            # Lorem ipsum placeholder for skills - matching exact format
            skills_text = """Design/UX Principles: Lorem ipsum, dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt.
AI & Tools: Lorem ipsum, dolor sit amet, consectetur adipiscing elit.
Frontend UX: Lorem ipsum, dolor sit amet, consectetur adipiscing elit.
Backend: Lorem ipsum, dolor sit amet, consectetur adipiscing elit.
Other: Lorem ipsum, dolor sit amet, consectetur adipiscing elit."""
        
        # Generate the full resume text
        name = resume_data['name'] if resume_data['name'] else "Lorem Ipsum"
        title = resume_data['title'] if resume_data['title'] else "DOLOR SIT AMET | CONSECTETUR ADIPISCING | ELIT SED DO"
        summary = resume_data['summary'] if resume_data['summary'] else "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur."
        
        resume_text = f"""
{name}
{contact_info}
{title}
{summary}

EXPERIENCE
{experience_text}
EDUCATION
{education_text}
TECHNICAL & CREATIVE SKILLS
{skills_text}
"""
        
        return resume_text
    
    def generate_pdf(self, resume_data: Dict, template: str) -> bytes:
        """Generate PDF resume"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=6,
            alignment=1,  # Center
            textColor=colors.black
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=self.styles['Heading2'],
            fontSize=12,
            spaceAfter=6,
            spaceBefore=12,
            textColor=colors.black
        )
        
        # Name
        story.append(Paragraph(resume_data['name'] or 'Your Name', title_style))
        
        # Title
        if resume_data['title']:
            story.append(Paragraph(resume_data['title'], self.styles['Normal']))
        
        # Contact info
        contact_parts = []
        if resume_data['contact']['email']:
            contact_parts.append(resume_data['contact']['email'])
        if resume_data['contact']['phone']:
            contact_parts.append(resume_data['contact']['phone'])
        if resume_data['contact']['location']:
            contact_parts.append(resume_data['contact']['location'])
        
        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), self.styles['Normal']))
        
        story.append(Spacer(1, 12))
        
        # Summary
        if resume_data['summary']:
            story.append(Paragraph("PROFESSIONAL SUMMARY", header_style))
            story.append(Paragraph(resume_data['summary'], self.styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Experience
        if resume_data['experience']:
            story.append(Paragraph("WORK EXPERIENCE", header_style))
            for exp in resume_data['experience']:
                date_range = ""
                if exp.get('start_date') and exp.get('end_date'):
                    date_range = f" ({exp['start_date']} - {exp['end_date']})"
                elif exp.get('start_date'):
                    date_range = f" ({exp['start_date']} - Present)"
                
                story.append(Paragraph(f"{exp.get('title', '')} — {exp.get('company', '')}{date_range}", self.styles['Normal']))
                if exp.get('description'):
                    story.append(Paragraph(exp['description'], self.styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Education
        if resume_data['education']:
            story.append(Paragraph("EDUCATION", header_style))
            for edu in resume_data['education']:
                year_str = f" ({edu.get('year', '')})" if edu.get('year') else ""
                story.append(Paragraph(f"{edu.get('degree', '')} — {edu.get('school', '')}{year_str}", self.styles['Normal']))
        
        # Skills
        if resume_data['skills']:
            story.append(Paragraph("SKILLS", header_style))
            story.append(Paragraph(", ".join(resume_data['skills']), self.styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_docx(self, resume_data: Dict, template: str) -> bytes:
        """Generate DOCX resume"""
        doc = Document()
        
        # Name
        doc.add_heading(resume_data['name'] or 'Your Name', level=0)
        
        # Title
        if resume_data['title']:
            doc.add_paragraph(resume_data['title'])
        
        # Contact info
        contact_parts = []
        if resume_data['contact']['email']:
            contact_parts.append(resume_data['contact']['email'])
        if resume_data['contact']['phone']:
            contact_parts.append(resume_data['contact']['phone'])
        if resume_data['contact']['location']:
            contact_parts.append(resume_data['contact']['location'])
        
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))
        
        # Summary
        if resume_data['summary']:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(resume_data['summary'])
        
        # Experience
        if resume_data['experience']:
            doc.add_heading("Work Experience", level=1)
            for exp in resume_data['experience']:
                date_range = ""
                if exp.get('start_date') and exp.get('end_date'):
                    date_range = f" ({exp['start_date']} - {exp['end_date']})"
                elif exp.get('start_date'):
                    date_range = f" ({exp['start_date']} - Present)"
                
                doc.add_paragraph(f"{exp.get('title', '')} — {exp.get('company', '')}{date_range}")
                if exp.get('description'):
                    doc.add_paragraph(exp['description'])
        
        # Education
        if resume_data['education']:
            doc.add_heading("Education", level=1)
            for edu in resume_data['education']:
                year_str = f" ({edu.get('year', '')})" if edu.get('year') else ""
                doc.add_paragraph(f"{edu.get('degree', '')} — {edu.get('school', '')}{year_str}")
        
        # Skills
        if resume_data['skills']:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(resume_data['skills']))
        
        # Save to bytes
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()
