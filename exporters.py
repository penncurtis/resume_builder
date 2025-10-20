import io
from typing import Dict
from docx import Document
from bs4 import BeautifulSoup

# --- PDF from HTML using WeasyPrint (best), with graceful fallback ---
def export_pdf_from_html(html: str) -> bytes:
    try:
        from weasyprint import HTML
        pdf_io = io.BytesIO()
        HTML(string=html).write_pdf(pdf_io)
        pdf_io.seek(0)
        return pdf_io.getvalue()
    except Exception:
        # ultra-simple fallback: plain text dump (still produces a PDF via reportlab)
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        textobject = c.beginText(36, 750)
        textobject.setFont("Times-Roman", 11)
        soup = BeautifulSoup(html, "html.parser")
        txt = soup.get_text("\n")
        for line in txt.splitlines():
            textobject.textLine(line[:95])
            if textobject.getY() < 36:
                c.drawText(textobject); c.showPage(); textobject = c.beginText(36, 750); textobject.setFont("Times-Roman", 11)
        c.drawText(textobject); c.showPage(); c.save()
        buffer.seek(0)
        return buffer.getvalue()

# --- DOCX builder with proper formatting to match preview ---
def export_docx_from_data(data: Dict, template_name: str) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    
    doc = Document()
    
    # Set default font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(10)
    
    # Add name as main heading
    name = data.get("name", "Your Name")
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(24)
    heading_run.font.bold = True
    
    # Add title
    if data.get("title"):
        title_para = doc.add_paragraph(data["title"])
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.runs[0]
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = None  # Use default color
    
    # Add contact info
    contact = data.get("contact", {})
    contact_parts = []
    if contact.get("email"): contact_parts.append(contact["email"])
    if contact.get("phone"): contact_parts.append(contact["phone"])
    if contact.get("location"): contact_parts.append(contact["location"])
    if contact.get("linkedin"): contact_parts.append(f"LinkedIn {contact['linkedin']}")
    if contact.get("github"): contact_parts.append(f"GitHub {contact['github']}")
    
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact_run = contact_para.runs[0]
        contact_run.font.size = Pt(10)
    
    # Add summary
    if data.get("summary"):
        doc.add_heading("SUMMARY", level=1)
        summary_para = doc.add_paragraph(data["summary"])
        summary_run = summary_para.runs[0]
        summary_run.font.size = Pt(10)
    
    # Add work experience with proper formatting
    if data.get("experience"):
        doc.add_heading("WORK EXPERIENCE", level=1)
        for job in data["experience"]:
            # Job title and company
            job_header = f"{job.get('title', '')} — {job.get('company', '')}"
            if job.get('location'):
                job_header += f" | {job['location']}"
            
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(job_header)
            job_run.font.size = Pt(11)
            job_run.font.bold = True
            
            # Dates
            if job.get('start_date') or job.get('end_date'):
                dates = f"{job.get('start_date', '')} – {job.get('end_date', '')}"
                date_para = doc.add_paragraph(dates)
                date_run = date_para.runs[0]
                date_run.font.size = Pt(10)
                date_run.font.italic = True
            
            # Bullet points
            for bullet in job.get("bullets", []):
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_run = bullet_para.runs[0]
                bullet_run.font.size = Pt(10)
            
            # Technologies
            if job.get("technologies"):
                tech_para = doc.add_paragraph(f"Tech: {', '.join(job['technologies'])}")
                tech_run = tech_para.runs[0]
                tech_run.font.size = Pt(9)
                tech_run.font.italic = True
    
    # Add education
    if data.get("education"):
        doc.add_heading("EDUCATION", level=1)
        for edu in data["education"]:
            # Degree and school
            edu_header = f"{edu.get('degree', '')} — {edu.get('school', '')}"
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(edu_header)
            edu_run.font.size = Pt(11)
            edu_run.font.bold = True
            
            # Dates and location
            if edu.get('location') or edu.get('start_date') or edu.get('end_date'):
                location = edu.get('location', '')
                dates = f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
                if location and dates:
                    meta = f"{location} | {dates}"
                elif location:
                    meta = location
                elif dates:
                    meta = dates
                else:
                    meta = ""
                
                if meta:
                    meta_para = doc.add_paragraph(meta)
                    meta_run = meta_para.runs[0]
                    meta_run.font.size = Pt(10)
                    meta_run.font.italic = True
            
            # Details
            for detail in edu.get("details", []):
                detail_para = doc.add_paragraph(detail, style="List Bullet")
                detail_run = detail_para.runs[0]
                detail_run.font.size = Pt(10)
    
    # Add skills
    if data.get("skills"):
        doc.add_heading("SKILLS", level=1)
        for category, skills in data["skills"].items():
            if skills:
                skill_text = f"{category.replace('_', ' ').title()}: {', '.join(skills)}"
                skill_para = doc.add_paragraph(skill_text)
                skill_run = skill_para.runs[0]
                skill_run.font.size = Pt(10)
    
    # Set section heading styles
    for heading in doc.paragraphs:
        if heading.style.name.startswith('Heading 1'):
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(10)
            heading_run.font.bold = True
            heading_run.font.color.rgb = None
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()